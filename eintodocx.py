from docx import Document
from docx.shared import Inches
from docx.shared import Pt,Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement, qn
import glob
import sys
import os
import re

PAR_STATE_ERROR = 0
PAR_STATE_NORMAL = 1
PAR_STATE_UNDERLINE = 2
PAR_STATE_BOLD = 3

def parseEin(file):
    try:
        with open(file,mode='rb') as f:
            first = f.read(1)
            parsedList = []
            if first == b';':
                try:
                    fileText = f.read().decode('cp856')
                    lines = fileText.splitlines()
                    header = [lines[0]]
                    lineIndex = 1
                    for line in lines[lineIndex:]:
                        lineIndex += 1
                        if line.startswith(';'):
                            header.append(line)
                        else:
                            break
                    headerParams = header[0].split(',')
                    headerDict = {'pageTop':headerParams[1],'pageBottom':headerParams[2],'pageSize':headerParams[3],'pageLeft':headerParams[6], 'pageRight':headerParams[7]}
                    for line in lines[lineIndex:]:
                        parsedLine = []
                        normalText = ""
                        dotCommandActive = line.startswith('.')
                        if dotCommandActive:
                            skipCount = 2
                            dotCommand = line[1]
                            if dotCommand.lower() == 'p': #pagebreak
                                pass
                            elif dotCommand.lower() == 'a': #add lines
                                skipCount+=1
                                pass
                            elif dotCommand.lower() == 'h': #header
                                pass
                            elif dotCommand.lower() == 'f': #footer
                                pass
                            else:
                                skipCount = 0
                            line = line[skipCount:]
                        lastRun = {'text':'', 'bold':False,'underline':False, 'centered':False, 'left':False}
                        re_S = re.compile(r'(  +)')
                        sline = re_S.split(line)
                        sequenceLen = 0
                        lastControlCharacter = False
                        for ll in sline:
                            lastRun = {'text':'', 'bold':False,'underline':False, 'centered':False, 'left':False}
                            for c in ll:
                                controlCharacter = True
                                if c == '╩': #Bold
                                    if not lastRun['underline'] and lastRun['text']!='':
                                            parsedLine.append(lastRun)
                                            lastRun = {'text':'', 'bold':False,'underline':False, 'centered':False, 'left':False}
                                    lastRun['bold'] = True
                                elif c == '╦': #Underline bold
                                    if (not lastRun['underline'] and not lastRun['bold']) and lastRun['text']!='':
                                            parsedLine.append(lastRun)
                                            lastRun = {'text':'', 'bold':False,'underline':False, 'centered':False, 'left':False}
                                    lastRun['bold'] = True
                                    lastRun['underline'] = True
                                elif c == '╔': #Underline
                                    if not lastRun['underline'] and lastRun['text']!='':
                                            parsedLine.append(lastRun)
                                            lastRun = {'text':'', 'bold':False,'underline':False, 'centered':False, 'left':False}
                                    lastRun['underline'] = True
                                elif c == '█': #Printer Command
                                    pass
                                elif c == '▄': #Printer Command 2
                                    pass
                                elif c == '┘': #Center Align
                                    lastRun['centered'] = True
                                elif c == '┌': #Left Aling
                                    lastRun['left'] = True
                                elif c == '\x18': #UpperMark
                                    pass
                                elif c == '\x19': #LowerMark
                                    pass
                                else:
                                    controlCharacter = False
                                    if False: #c=='.' or c==',' or c=='[' or c==']' or c=='(' or c==')':
                                        if lastRun['text']!='':
                                            parsedLine.append(lastRun)
                                        lastRun = {'text':c, 'bold':False,'underline':False, 'centered':False, 'left':False}
                                        parsedLine.append(lastRun)
                                        lastRun = {'text':'', 'bold':False,'underline':False, 'centered':False, 'left':False}
                                        sequenceLen = 0
                                    elif False: #c==' ':
                                        if sequenceLen:
                                            parsedLine.append(lastRun)
                                            lastRun = {'text':'', 'bold':False,'underline':False, 'centered':False, 'left':False}
                                            sequenceLen = 0
                                        lastRun['text'] += ' '
                                    else:
                                        '''
                                        if lastRun['text']!='' and not sequenceLen:
                                            parsedLine.append(lastRun)
                                            lastRun = {'text':'', 'bold':False,'underline':False, 'centered':False, 'left':False}
                                        sequenceLen += 1
                                        '''
                                        if not lastControlCharacter:
                                            if lastRun['bold'] or lastRun['underline']:
                                                if lastRun['text'].strip()!='':
                                                    parsedLine.append(lastRun) 
                                                    lastRun = {'text':'', 'bold':False,'underline':False, 'centered':False, 'left':False}
                                        lastRun['text'] += c
                                        
                                lastControlCharacter = controlCharacter
                            parsedLine.append(lastRun)    
                        parsedList.append(parsedLine)
                    saveDocx(parsedList, file)
                    return
                except Exception as e:
                    print(e)
                    return
    except Exception as e:
        print(e)
    return

def saveDocx(parsedList, file):
    document = Document()
    document.sections[-1].left_margin = Pt(64)
    document.sections[-1].right_margin = Pt(64)
    document.sections[-1].top_margin = Cm(1.4)
    document.sections[-1].bottom_margin = Cm(1.4)
    normalStype = document.styles['Normal']
    normalStype.font.name = 'Courier New'
    normalStype.font.size = Pt(10)
    #document.add_heading('Document Title', 0)
    lineIndex = 0
    for e in parsedList:
        lineIndex += 1
        p = document.add_paragraph('')
        p.style = normalStype
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after = 0
        p.paragraph_format.line_spacing = Pt(11)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        ppr = p._element.get_or_add_pPr()
            
        w_nsmap = '{'+ppr.nsmap['w']+'}'
        bidi = None
        jc = None
        cs = None
        for element in ppr:
            if element.tag == w_nsmap + 'bidi':
                bidi = element
        if bidi is None:
            bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'),'he-IL')
        ppr.append(bidi)
        for run in e:
            runText = run['text']
            if runText.startswith('-') and runText.endswith('-') and runText[1:-1].isnumeric():
                footer = document.sections[-1].footer
                footer.is_linked_to_previous = False
                fp = footer.paragraphs[0]
                fp.add_run(runText)
                fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                document.add_section()
            else:
                runText = runText + '\u200F'    
                r = p.add_run(runText)
                if lineIndex > 60 and len(document.sections)<2:
                    document.add_section()
                r.cs_size = Pt(10)
                r.cs_bold = run['bold']
                r.bold = run['bold']
                r.underline = run['underline']
                if run['centered']:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if run['left']:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                ppr = r._element.get_or_add_rPr()
                    
                w_nsmap = '{'+ppr.nsmap['w']+'}'
                bidi = None
                for element in ppr:
                    if element.tag == w_nsmap + 'bidi':
                        bidi = element
                if bidi is None:
                    bidi = OxmlElement('w:bidi')
                bidi.set(qn('w:val'),'he-IL')
                ppr.append(bidi)
                
    '''
    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    #document.add_picture('monty-truth.png', width=Inches(1.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    '''
    document.save('out/demo-'+ os.path.split(file)[-1] +'.docx')

if __name__ == "__main__":
    if len(sys.argv) < 2:
        exit(0)
    for arg in sys.argv[1:]:
        for file in glob.glob(arg):
            parseEin(file)
    exit(0)